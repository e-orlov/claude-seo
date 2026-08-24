"""
report_renderer.py
══════════════════
Generic renderer for SEO audit reports.

This file is static infrastructure — it contains no audit logic, no domain knowledge,
no column definitions, no issue types. It renders whatever CFG and SECTIONS contain.

Usage in every audit report script:

    from report_renderer import render

    CFG = { ... }       # audit metadata
    SECTIONS = [ ... ]  # derived by Claude from analysis — structure and content

    render(CFG, SECTIONS)

---

SECTIONS block types
────────────────────

{"type": "heading", "text": "...", "level": 1}
    Section heading. level=1: page break before, NAVY 13pt.
                     level=2: no page break, DBLUE 11pt.

{"type": "bt", "text": "..."}
    Body text paragraph. 9.5pt BLACK.

{"type": "note", "text": "..."}
    Italic caption / note. 8.5pt MGREY.

{"type": "spacer"}
    Empty paragraph for vertical spacing.

{"type": "table",
 "columns": ["Col A", "Col B", ...],   # header labels — derived from analysis
 "widths":  [5.0, 8.5, ...],           # cm — must sum to ~26.1 cm
 "rows": [
     {"bg": "LYELL", "cells": ["val1", "val2", ...]},
     {"bg": "WHITE", "cells": ["val3", "val4", ...]},
     ...
 ],
 "summary": {"bg": "NAVY", "cells": ["Total", "...", ...]},  # optional footer row
}

    Column count must match across columns, widths, every row's cells, and summary.
    bg values: "LRED" | "LYELL" | "LGREEN" | "LGREY" | "MGREY" | "WHITE" | "NAVY" | any hex string
    All cell values must be pre-formatted strings. The renderer does no calculations.
    pct(), date formatting, number formatting — all done by Claude before writing SECTIONS.

---

CFG required keys
─────────────────
    domain        str   e.g. "example.com"
    report_title  str   e.g. "Interne Weiterleitungen"
    report_name   str   filename stem, e.g. "Interne_Weiterleitungen"
    date_label    str   e.g. "Juni 2026"
    date_slug     str   e.g. "2026-06"
    data_sources  str   e.g. "Screaming Frog MCP, Crawl Juni 2026"
    subtitle      str   shown below title on page 1
    baseline      int   indexable HTML URL count — from file_inventory
    output_dir    None  or explicit path string — None triggers auto-resolve

---

Column width constraint
───────────────────────
Total usable width = 29.7 - 1.8 - 1.8 = 26.1 cm
Sum of widths per table should be ≈ 26.1 cm (minor rounding acceptable).
"""

from docx.enum.table import WD_TABLE_ALIGNMENT

from docx_helpers import (
    setup_document,
    h, bt, note,
    hdr, cp, set_bg, scw,
    NAVY, LGREY, WHITE, BLACK,
    LYELL, LGREEN, LRED, MGREY,
)
from report_config import resolve_output_path


# ── Color lookup ──────────────────────────────────────────────────────────────

_BG_MAP = {
    "NAVY":   NAVY,
    "LGREY":  LGREY,
    "WHITE":  WHITE,
    "LYELL":  LYELL,
    "LGREEN": LGREEN,
    "LRED":   LRED,
    "MGREY":  MGREY,
    "BLACK":  BLACK,
}


def _resolve_bg(value: str) -> str:
    """
    Resolve a bg value to a hex string.
    Accepts named keys ("LRED", "LYELL", ...) or bare 6-char hex strings ("FCE4D6").
    """
    if value in _BG_MAP:
        return _BG_MAP[value]
    if len(value) == 6 and all(c in "0123456789ABCDEFabcdef" for c in value):
        return value
    raise ValueError(
        f"report_renderer: unknown bg value '{value}'. "
        f"Use one of {list(_BG_MAP)} or a 6-char hex string."
    )


# ── Table renderer ────────────────────────────────────────────────────────────

def _render_table(doc, block: dict, cfg: dict) -> None:
    """
    Render a single table block.

    Validates column count consistency across columns, widths, rows and summary.
    Raises ValueError on mismatch so errors surface at render time, not silently.
    """
    columns  = block["columns"]
    widths   = block["widths"]
    rows     = block.get("rows", [])
    summary  = block.get("summary")
    n_cols   = len(columns)

    # Validate widths
    if len(widths) != n_cols:
        raise ValueError(
            f"report_renderer: table has {n_cols} columns but {len(widths)} widths. "
            f"columns={columns}"
        )

    # Validate rows
    for ridx, row in enumerate(rows):
        if len(row["cells"]) != n_cols:
            raise ValueError(
                f"report_renderer: row {ridx} has {len(row['cells'])} cells "
                f"but table has {n_cols} columns. "
                f"columns={columns}, cells={row['cells']}"
            )

    # Validate summary
    if summary and len(summary["cells"]) != n_cols:
        raise ValueError(
            f"report_renderer: summary has {len(summary['cells'])} cells "
            f"but table has {n_cols} columns."
        )

    # Build table
    tbl = doc.add_table(rows=1, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr(tbl, columns, widths)

    # Data rows
    for row in rows:
        bg  = _resolve_bg(row.get("bg", "WHITE"))
        tr  = tbl.add_row()
        for i, cell_value in enumerate(row["cells"]):
            cp(tr.cells[i], cell_value, fs=block.get("font_size", 9))
            set_bg(tr.cells[i], bg)

    # Optional summary row (NAVY background, white text)
    if summary:
        bg  = _resolve_bg(summary.get("bg", "NAVY"))
        tr  = tbl.add_row()
        for i, cell_value in enumerate(summary["cells"]):
            col = WHITE if bg == NAVY else BLACK
            cp(tr.cells[i], cell_value, bold=True, fs=block.get("font_size", 9),
               align="center", col=col)
            set_bg(tr.cells[i], bg)

    # Re-apply column widths (docx requires this after all rows are added)
    scw(tbl, widths)


# ── Main render entry point ───────────────────────────────────────────────────

def render(cfg: dict, sections: list) -> None:
    """
    Render a complete report document from cfg and sections.

    Args:
        cfg:      ReportConfig dict — see module docstring for required keys.
        sections: List of block dicts — see module docstring for supported types.

    Saves the document to the path resolved by resolve_output_path(cfg) and
    prints the saved path to stdout.
    """
    # Validate required CFG keys
    required = ("domain", "report_title", "report_name",
                "date_label", "date_slug", "data_sources", "subtitle", "baseline")
    missing = [k for k in required if not cfg.get(k) and cfg.get(k) != 0]
    if missing:
        raise ValueError(
            f"report_renderer: CFG is missing required keys: {missing}"
        )

    doc, sec = setup_document(cfg)

    for block in sections:
        block_type = block.get("type")

        if block_type == "heading":
            h(doc, block["text"], lv=block.get("level", 1))

        elif block_type == "bt":
            bt(doc, block["text"])

        elif block_type == "note":
            note(doc, block["text"])

        elif block_type == "spacer":
            doc.add_paragraph()

        elif block_type == "table":
            _render_table(doc, block, cfg)
            doc.add_paragraph()

        else:
            raise ValueError(
                f"report_renderer: unknown block type '{block_type}'. "
                f"Supported: heading, bt, note, spacer, table."
            )

    output_path = resolve_output_path(cfg)
    doc.save(output_path)
    print("Saved:", output_path)
