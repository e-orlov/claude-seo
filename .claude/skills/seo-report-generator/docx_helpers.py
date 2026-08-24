"""
docx_helpers.py
═══════════════
Shared helper library for all SEO report scripts.

Every gen_*.py imports from here. No client-specific data allowed in this file.
All colors, helpers, and document setup are defined once.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Color palette ─────────────────────────────────────────────────────────────
# Single source of truth. All scripts use these hex strings.
# Severity mapping:
#   Kritisch  → LRED
#   Hoch      → LYELL
#   Mittel    → LGREEN
#   Niedrig   → LGREY
#   Positiv   → LGREEN
#   Neutral   → WHITE or LGREY (alternating)

NAVY   = "1F3864"   # table headers, primary headings
DBLUE  = "2E5F8A"   # lv=2 sub-headings
LGREY  = "F2F2F2"   # alternating row background / Niedrig
WHITE  = "FFFFFF"   # default row background
BLACK  = "000000"   # body text
LYELL  = "FFF2CC"   # Hoch / warning
LGREEN = "E2EFDA"   # Mittel / ok
LRED   = "FCE4D6"   # Kritisch / error
MGREY  = "595959"   # footer, captions, notes

# Priority → background color mapping (keys match both German casing variants)
PRIO_BG = {
    "Kritisch": LRED,   "KRITISCH": LRED,
    "Hoch":     LYELL,  "HOCH":     LYELL,
    "Mittel":   LGREEN, "MITTEL":   LGREEN,
    "Niedrig":  LGREY,  "NIEDRIG":  LGREY,
}


# ── Core helpers ──────────────────────────────────────────────────────────────

def hex_rgb(h: str) -> RGBColor:
    """Convert a 6-char hex string to RGBColor."""
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def pct(n: int, baseline: int, label: str = "URLs") -> str:
    """
    Format an absolute count with its percentage of the baseline.

    Args:
        n:        affected count
        baseline: total reference count (from CFG["baseline"])
        label:    unit label, default "URLs"

    Returns:
        e.g. "1.180 (24,2 % von 4.866 URLs)"

    Note: use pct() only for page/URL counts.
    Do not use for link counts, kB values, image counts or other non-URL metrics.
    """
    if baseline == 0:
        return str(n)
    return f"{n} ({n / baseline * 100:.1f} % von {baseline} {label})"


def stripe(i: int) -> str:
    """Alternating row background: even rows LGREY, odd rows WHITE."""
    return LGREY if i % 2 == 0 else WHITE


# ── Cell helpers ──────────────────────────────────────────────────────────────

def set_bg(cell, hx: str) -> None:
    """Set cell background to a hex color string."""
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), hx)
    e = pr.find(qn("w:shd"))
    if e is not None:
        pr.remove(e)
    pr.append(s)


def cp(cell, text, bold=False, fs=9, align="left", col=None) -> None:
    """
    Write text into a cell paragraph.

    Args:
        cell:  docx table cell
        text:  cell content (will be str-converted)
        bold:  bold run
        fs:    font size in pt
        align: "left" | "center" | "right"
        col:   hex color string; defaults to BLACK
    """
    if col is None:
        col = BLACK
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.clear()
    al = {
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    p.alignment = al
    pp = p._p.get_or_add_pPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), "20")
    sp.set(qn("w:after"), "20")
    pp.append(sp)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(fs)
    r.font.color.rgb = hex_rgb(col)


def scw(table, wds: list) -> None:
    """Set column widths in cm for all rows of a table."""
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(wds):
                tc = cell._tc
                pr = tc.get_or_add_tcPr()
                w = OxmlElement("w:tcW")
                w.set(qn("w:w"), str(int(wds[i] * 567)))
                w.set(qn("w:type"), "dxa")
                e = pr.find(qn("w:tcW"))
                if e is not None:
                    pr.remove(e)
                pr.append(w)


def hdr(tbl, labels: list, wds: list) -> None:
    """
    Write header row into table row 0 and set column widths.

    Header cells: white bold text on NAVY background, centered.
    """
    row = tbl.rows[0]
    for i, label in enumerate(labels):
        cp(row.cells[i], label, bold=True, fs=9, align="center", col=WHITE)
        set_bg(row.cells[i], NAVY)
    scw(tbl, wds)


# ── Paragraph helpers ─────────────────────────────────────────────────────────

def h(doc, text: str, lv: int = 1):
    """
    Add a section heading paragraph.

    Args:
        doc:  Document object
        text: heading text
        lv:   1 = main section (NAVY, 13pt, always page break before)
              2 = sub-heading (DBLUE, 11pt, no page break)
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    if lv == 1:
        pPr = p._p.get_or_add_pPr()
        pbB = OxmlElement("w:pageBreakBefore")
        pbB.set(qn("w:val"), "1")
        pPr.append(pbB)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13 if lv == 1 else 11)
    r.font.color.rgb = hex_rgb(NAVY if lv == 1 else DBLUE)
    return p


def bt(doc, text: str):
    """Add a body text paragraph (9.5pt, BLACK)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(9.5)
    r.font.color.rgb = hex_rgb(BLACK)
    return p


def note(doc, text: str):
    """Add an italic note/caption paragraph (8.5pt, MGREY)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = hex_rgb(MGREY)
    return p


# ── Document setup ────────────────────────────────────────────────────────────

def setup_document(cfg: dict):
    """
    Create and configure a landscape A4 Word document.

    Args:
        cfg: ReportConfig dict. Required keys:
               domain       — e.g. "example.com"
               report_title — e.g. "Interne Weiterleitungen"
               date_label   — e.g. "Juni 2026"
               data_sources — e.g. "Screaming Frog MCP, Crawl Juni 2026"
               subtitle     — subtitle line shown below main title

    Returns:
        (doc, sec) — Document and first section objects

    Page layout:  29.7 × 21.0 cm, margins 1.8 / 1.8 / 1.5 / 1.5 cm
    Font:         Calibri applied to all styles
    Footer:       "{domain} -- {report_title} | Datenbasis: {data_sources} | Stand: {date_label}"
                  No Claude/tool attribution.
    Title block:  centered title + italic subtitle — NO page break before or after.
    """
    doc = Document()
    sec = doc.sections[0]

    # Page: landscape A4
    sec.page_width    = Cm(29.7)
    sec.page_height   = Cm(21.0)
    sec.left_margin   = Cm(1.8)
    sec.right_margin  = Cm(1.8)
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.5)

    # Font: Calibri across all styles
    for s in doc.styles:
        if hasattr(s, "font"):
            try:
                s.font.name = "Calibri"
            except Exception:
                pass

    # Footer
    ftr = sec.footer
    fp  = ftr.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text = (
        f"{cfg['domain']} -- {cfg['report_title']}"
        f" | Datenbasis: {cfg['data_sources']}"
        f" | Stand: {cfg['date_label']}"
    )
    fr = fp.add_run(footer_text)
    fr.font.size = Pt(8)
    fr.font.color.rgb = hex_rgb(MGREY)

    # Title block (no page break before or after)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{cfg['report_title']} -- {cfg['domain']}")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = hex_rgb(NAVY)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(cfg.get("subtitle", ""))
    r2.font.size = Pt(9)
    r2.font.color.rgb = hex_rgb(MGREY)
    r2.italic = True

    doc.add_paragraph()  # spacing after title block

    return doc, sec
